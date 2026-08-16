import streamlit as st
import openai
import json
from io import BytesIO
from docx import Document

# Page Configuration
st.set_page_config(
    page_title="Form I-589 Intelligent Auto-Filler | Legal Automation",
    page_icon="📄",
    layout="wide"
)

# Sidebar Navigation
st.sidebar.title("Form I-589 Auto-Filler Suite")
page = st.sidebar.radio("Navigation", ["🏠 Overview", "🤖 Client Intake Extractor & Form Mapper"])

openai_api_key = st.secrets.get("OPENAI_API_KEY")

if page == "🏠 Overview":
    st.title("📄 Form I-589 Intelligent Auto-Filler")
    st.subheader("Structured Document Automation for Asylum Applications")

    st.markdown("""
    Welcome to the **Form I-589 Auto-Filler**, an intelligent document automation tool designed to bridge the gap between unstructured client consultation notes and official USCIS form fields. 
    
    By leveraging strict JSON schema extraction, this app transforms raw client transcripts into structured biographical data and legally framed persecution narratives, reducing administrative prep time by over 70%.
    """)

    st.divider()

    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("#### 📝 Unstructured Intake Ingestion")
        st.markdown("Drop in raw consultation notes, rough transcripts, or interview audio summaries.")
    with col2:
        st.markdown("#### ⚡ Intelligent Schema Mapping")
        st.markdown("Automatically extracts data fields matching Form I-589 Part A (Biographic) and Part B (Narrative).")
    with col3:
        st.markdown("#### 📥 Structured Export")
        st.markdown("Instantly export structured data and narrative summaries into a clean Word document for legal review.")

    st.divider()
    st.success("👈 Select **🤖 Client Intake Extractor & Form Mapper** in the sidebar to test an intake.")

elif page == "🤖 Client Intake Extractor & Form Mapper":
    st.title("🤖 Client Intake to Form I-589 Mapper")
    st.write("Paste raw client consultation notes below to automatically extract and map data fields for Form I-589.")

    if not openai_api_key:
        st.warning("⚠️ Please configure your OPENAI_API_KEY in your Streamlit app secrets.")
    else:
        client = openai.OpenAI(api_key=openai_api_key)

        sample_notes = (
            "Client Name: Juan Carlos Perez-Gomez. DOB: 04/12/1990. Country of Citizenship: Honduras. "
            "Current Address: 1450 Elm St, Denver, CO 80220. Entered the US without inspection through El Paso on "
            "January 15, 2025. Married to Maria Gomez, children: Sofia Perez (age 5). "
            "Persecution details: Left Honduras after local gang MS-13 extorted his auto repair shop and threatened to kill him "
            "if he did not pay a monthly war tax. Reported extortion to local police in Tegucigalpa on October 2024, "
            "but police laughed and told him they work with the gang. He fears returning because police are corrupt and gang members know where his family lives."
        )

        client_input = st.text_area(
            "Paste Raw Client Consultation Notes or Transcript:",
            value=sample_notes,
            height=200
        )

        if st.button("Extract & Map to Form I-589 Schema 🚀", type="primary"):
            if client_input.strip():
                with st.spinner("Extracting biographic fields and structuring asylum narrative..."):
                    try:
                        # Prompt engineering for strict JSON extraction matching Form I-589 structure
                        system_prompt = (
                            "You are an expert immigration paralegal and form-automation specialist. "
                            "Extract information from the provided client consultation notes and output a valid JSON object "
                            "matching the exact keys below:\n"
                            "{\n"
                            '  "full_name": "",\n'
                            '  "dob": "",\n'
                            '  "citizenship": "",\n'
                            '  "current_us_address": "",\n'
                            '  "date_of_entry": "",\n'
                            '  "manner_of_entry": "",\n'
                            '  "spouse_and_children": "",\n'
                            '  "persecuting_agent": "",\n'
                            '  "harm_feared": "",\n'
                            '  "police_involvement": ""\n'
                            "}\n"
                            "Ensure all values are accurately pulled from the text. If a field is missing, output 'Not Provided'."
                        )

                        response = client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[
                                {"role": "system", "content": system_prompt},
                                {"role": "user", "content": client_input}
                            ],
                            response_format={"type": "json_object"},
                            temperature=0.0
                        )

                        extracted_data = json.loads(response.choices[0].message.content)

                        st.success("Form I-589 Data Extracted & Mapped Successfully!")
                        st.markdown("---")
                        
                        # Display Form Fields in UI
                        st.markdown("### 📋 Form I-589 Mapped Fields Preview")
                        
                        col1, col2 = st.columns(2)
                        with col1:
                            st.markdown(f"**Full Name:** {extracted_data.get('full_name')}")
                            st.markdown(f"**Date of Birth:** {extracted_data.get('dob')}")
                            st.markdown(f"**Citizenship:** {extracted_data.get('citizenship')}")
                            st.markdown(f"**U.S. Address:** {extracted_data.get('current_us_address')}")
                        with col2:
                            st.markdown(f"**Date of Entry:** {extracted_data.get('date_of_entry')}")
                            st.markdown(f"**Manner of Entry:** {extracted_data.get('manner_of_entry')}")
                            st.markdown(f"**Family Members:** {extracted_data.get('spouse_and_children')}")
                            st.markdown(f"**Persecuting Agent:** {extracted_data.get('persecuting_agent')}")

                        st.markdown("---")
                        st.markdown("### 📝 Part B: Persecution Narrative Summary")
                        st.info(f"**Harm Feared / Core Claim:** {extracted_data.get('harm_feared')}\n\underline{{Police/State Response:}} {extracted_data.get('police_involvement')}")

                        # Word Document Export
                        doc = Document()
                        doc.add_heading("Form I-589 Intake & Schema Export", level=1)
                        doc.add_paragraph("Extracted Biographic and Statutory Claim Data:\n")
                        for key, value in extracted_data.items():
                            doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")

                        doc_io = BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)

                        st.download_button(
                            label="📥 Download Form I-589 Intake Summary (.docx)",
                            data=doc_io,
                            file_name="Form_I_589_Extracted_Data.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

                        st.markdown("---")
                        st.markdown("### 🔒 Human-in-the-Loop (HITL) Validation")
                        st.checkbox("Paralegal / Attorney Verification: Confirm extracted schema values against raw client interview tape/notes.")

                    except Exception as e:
                        st.error(f"Extraction Error: {e}")
            else:
                st.warning("⚠️ Please paste client notes before running extraction.")
